from docx import Document
from docx.shared import Inches
from io import BytesIO
import matplotlib.pyplot as plt

def save_report_to_word(title, input_tables=None, output_tables=None, texts=None, images=None):
    doc = Document()
    doc.add_heading(title, level=1)

    # Вводные таблицы
    if input_tables:
        for name, df in input_tables.items():
            doc.add_heading(name, level=2)
            table = doc.add_table(rows=(df.shape[0] + 1), cols=df.shape[1])
            table.style = 'Table Grid'
            for j, col in enumerate(df.columns):
                table.cell(0, j).text = str(col)
            for i in range(df.shape[0]):
                for j in range(df.shape[1]):
                    table.cell(i + 1, j).text = str(df.iat[i, j])

    # Выводные таблицы
    if output_tables:
        for name, df in output_tables.items():
            doc.add_heading(name, level=2)
            table = doc.add_table(rows=(df.shape[0] + 1), cols=df.shape[1])
            table.style = 'Table Grid'
            for j, col in enumerate(df.columns):
                table.cell(0, j).text = str(col)
            for i in range(df.shape[0]):
                for j in range(df.shape[1]):
                    table.cell(i + 1, j).text = str(df.iat[i, j])

    # Тексты
    if texts:
        for txt in texts:
            doc.add_paragraph(txt)

    # Изображения
    if images:
        for fig in images:
            img_stream = BytesIO()
            fig.savefig(img_stream, format='png', bbox_inches='tight')
            img_stream.seek(0)
            doc.add_picture(img_stream, width=Inches(6))
            doc.add_paragraph()  # Отступ между картинками

    # Сохранение в байтовый поток
    final_buffer = BytesIO()
    doc.save(final_buffer)
    final_buffer.seek(0)
    return final_buffer
